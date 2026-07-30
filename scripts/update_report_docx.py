"""Synchronise le rapport Word avec les preuves finales gouvernées.

CE SCRIPT EST HORS SERVICE, ET C'EST ÉCRIT ICI PLUTÔT QUE DÉCOUVERT À
L'EXÉCUTION.

Trois constats établis par la lecture intégrale de ce fichier :

1. `REPORT` désigne `reports/Rapport_technique_E7301.docx`, **qui n'existe
   pas**. Le dépôt porte deux `.docx` à la racine, aux noms différents. Toute
   exécution lève donc `FileNotFoundError` avant la première ligne utile.

2. La table `replacements` est une liste de substitutions littérales figées
   — « 28 points d'entrée » → « 42 routes /api/ », « 17 features » →
   « 10 features contractuelles ». Ces valeurs sont des instantanés d'un état
   passé : l'API expose aujourd'hui **46** routes, pas 42. Réparer le chemin
   sans revoir la table ferait donc écrire des chiffres FAUX dans le rapport,
   ce qui est pire que de ne rien écrire.

3. `metrics["tests"]["tests"]` relit `project_metrics.json`, dont l'artefact
   courant annonce 153 tests alors que la suite en compte 262.

La substitution littérale sur un document qui bouge est le problème de fond,
pas le chemin. Le sort de ce script — réécriture ou suppression — relève d'un
arbitrage de périmètre porté au plan final ; il ne doit pas être exécuté d'ici
là, et son échec immédiat garantit qu'il ne le sera pas par mégarde.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

REPORT = Path("reports/Rapport_technique_E7301.docx")
METRICS = Path("reports/project_metrics.json")
ADDENDUM_TITLE = "Addendum final gouverné — 25 juillet 2026"


def all_paragraphs(document: Document):
    """Itère sur les paragraphes du corps, tableaux, en-têtes et pieds."""
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def replace_text(document: Document, replacements: dict[str, str]) -> None:
    """Remplace même lorsque Word a fragmenté une phrase sur plusieurs runs."""
    for paragraph in all_paragraphs(document):
        original = paragraph.text
        updated = original
        for old, new in replacements.items():
            updated = updated.replace(old, new)
        if updated == original:
            continue
        if paragraph.runs:
            paragraph.runs[0].text = updated
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.text = updated


def set_cell(cell, value: str, *, bold: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.text = ""
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.size = Pt(9)


def style_tables(document: Document) -> None:
    """Répète les en-têtes et empêche la fragmentation des lignes."""
    for table in document.tables:
        if table.rows:
            tr_pr = table.rows[0]._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:tblHeader")) is None:
                header = OxmlElement("w:tblHeader")
                header.set(qn("w:val"), "true")
                tr_pr.append(header)
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:cantSplit")) is None:
                tr_pr.append(OxmlElement("w:cantSplit"))


def synchronize_model_tables(document: Document, metrics: dict) -> None:
    """Aligne les tableaux du modèle sur le contrat runtime exact."""
    test_count = metrics["tests"]["tests"]
    feature_rows = [
        ("duty_residual_z", "Résidu thermique standardisé"),
        ("duty_residual_trend_14d", "Tendance causale 14 jours du résidu"),
        ("conc_min", "Titre acide minimum des deux analyseurs"),
        ("conc_bias_drift_z", "Dérive standardisée du biais de titre"),
        ("conc_drop_24h", "Variation causale du titre sur 24 h"),
        ("flow_per_load", "Débit acide normalisé par la charge"),
        ("d_t_out", "Variation horaire de la température de sortie"),
        ("d_conc", "Variation horaire du titre acide"),
        ("t_out_local_z", "Déviation locale de la température de sortie"),
        ("t_in_local_z", "Déviation locale de la température d'entrée"),
    ]
    for table in document.tables:
        rows_text = [
            [cell.text.strip() for cell in row.cells]
            for row in table.rows
        ]
        flattened = " | ".join(value for row in rows_text for value in row)
        for row in table.rows:
            if (
                row.cells
                and row.cells[0].text.strip() == "Tests automatisés"
                and len(row.cells) >= 2
            ):
                set_cell(row.cells[1], str(test_count))
            if not row.cells or len(row.cells) < 2:
                continue
            label = row.cells[0].text.strip()
            if label == "Tests logiciels":
                tests = metrics["tests"]
                set_cell(
                    row.cells[1],
                    f"{tests['tests']} réussis, {tests['failures']} échec, "
                    f"{tests['errors']} erreur, {tests['skipped']} ignoré",
                )
            elif label == "Couverture":
                coverage = metrics["coverage"]
                set_cell(
                    row.cells[1],
                    f"{coverage['percent']:.2f} % "
                    f"({coverage['covered_lines']}/{coverage['statements']} lignes)",
                )
            elif label == "Routes API":
                set_cell(row.cells[1], str(metrics["api"]["route_count"]))
        if "Isolation Forest" in flattened and "Seuil de décision" in flattened:
            updates = {
                "Features": "10 features contractuelles ordonnées",
                "Période d’apprentissage": (
                    "06/01/2024 → 19/07/2024 (3 294 observations)"
                ),
                "Seuil de décision": "0,973",
            }
            for row in table.rows:
                if row.cells[0].text.strip() in updates and len(row.cells) >= 2:
                    set_cell(row.cells[1], updates[row.cells[0].text.strip()])
        if (
            rows_text
            and rows_text[0][:2] == ["Feature", "Nature"]
            and "delta_t" in flattened
        ):
            while len(table.rows) < len(feature_rows) + 1:
                table.add_row()
            set_cell(table.rows[0].cells[0], "Feature", bold=True)
            set_cell(table.rows[0].cells[1], "Nature", bold=True)
            for index, (feature, nature) in enumerate(feature_rows, 1):
                set_cell(table.rows[index].cells[0], feature)
                set_cell(table.rows[index].cells[1], nature)
            for row in list(table.rows[len(feature_rows) + 1 :]):
                table._tbl.remove(row._tr)


def add_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    # Le modèle Word source utilise des noms de styles localisés. Réutiliser le
    # premier style de tableau évite toute dépendance à "Table Grid".
    if document.tables[:-1] and document.tables[0].style is not None:
        table.style = document.tables[0].style
    set_cell(table.rows[0].cells[0], "Preuve", bold=True)
    set_cell(table.rows[0].cells[1], "Valeur vérifiée", bold=True)
    for label, value in rows:
        cells = table.add_row().cells
        set_cell(cells[0], label, bold=True)
        set_cell(cells[1], value)


def append_addendum(document: Document, metrics: dict) -> None:
    if any(p.text.strip() == ADDENDUM_TITLE for p in document.paragraphs):
        return

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    title = document.add_heading(ADDENDUM_TITLE, level=1)
    title.runs[0].font.color.rgb = RGBColor(0x00, 0x62, 0x4B)

    document.add_paragraph(
        "Cet addendum prévaut sur toute valeur historique divergente du corps du "
        "rapport. Il est généré à partir de reports/project_metrics.json après "
        "exécution de la campagne automatisée finale."
    )

    document.add_heading("Périmètre factuel", level=2)
    data = metrics["data"]
    model = metrics["model"]
    tests = metrics["tests"]
    coverage = metrics["coverage"]
    add_table(
        document,
        [
            ("Équipement", "S-PC-E7301 — refroidisseur de séchage PS III"),
            (
                "Données",
                f"{data['raw_rows']} lignes brutes, "
                f"{data['usable_timestamps']} horodatages utilisables, "
                f"{data['dcs_tags']} tags",
            ),
            ("Période", f"{data['period'][0]} → {data['period'][1]}"),
            (
                "Déduplication",
                f"{data['duplicates_merged_with_audit']} horodatages fusionnés "
                "avec événements qualité explicites",
            ),
            ("Empreinte DATA.xlsx", data["sha256"]),
            (
                "Modèle runtime",
                f"{model['runtime_source']} — statut artefact "
                f"{model['artifact_promotion_status']}",
            ),
            (
                "Signal historique",
                f"{model['alert_hours_historical']} heures au-dessus du seuil, "
                f"{model['episodes']} épisodes candidats",
            ),
            (
                "Tests logiciels",
                f"{tests['tests']} réussis, {tests['failures']} échec, "
                f"{tests['errors']} erreur, {tests['skipped']} ignoré",
            ),
            (
                "Couverture",
                f"{coverage['percent']:.2f} % "
                f"({coverage['covered_lines']}/{coverage['statements']} lignes)",
            ),
            ("Routes API", str(metrics["api"]["route_count"])),
        ],
    )

    document.add_heading("Gouvernance du modèle", level=2)
    document.add_paragraph(
        "Le détecteur est une surveillance comportementale non supervisée. "
        "Il ne démontre ni une panne future ni une causalité mécanique. "
        "L'artefact candidat est refusé au runtime ; la démonstration utilise "
        "un modèle reconstruit localement et explicitement non promu."
    )
    add_table(
        document,
        [
            ("Features contractuelles ordonnées", ", ".join(model["ordered_features"])),
            ("Seuil du runtime", f"{model['threshold']:.6f}"),
            (
                "Gates obligatoires non satisfaits",
                ", ".join(model["artifact_failed_gates"]),
            ),
            (
                "Validation industrielle",
                "Non démontrée par les tests logiciels",
            ),
        ],
    )

    document.add_heading("Décision de déploiement", level=2)
    add_table(
        document,
        [
            ("Démonstration hors ligne", "GO"),
            (
                "Pilote silencieux",
                "GO conditionnel — aucune commande procédé, qualification terrain",
            ),
            ("Pilote assisté", "NO-GO en l'état"),
            ("Production", "NO-GO en l'état"),
        ],
    )
    p = document.add_paragraph()
    run = p.add_run(
        "Limites bloquantes : absence de labels GMAO, absence de validation "
        "externe/hors site, stabilité hors période non acquise, plans mécaniques "
        "cotés indisponibles, connecteurs OIDC/SMTP/CMMS non qualifiés et absence "
        "de tests utilisateurs opérateurs."
    )
    run.bold = True

    document.add_heading("Règles d'interprétation", level=2)
    for text in (
        "Le modèle 3D est conceptuel et non dimensionnel ; aucune quantité de tubes "
        "n'est revendiquée sans plans 711-104/105/106.",
        "Le Judge est un contrôleur de cohérence interne indépendant du texte "
        "généré, mais pas de la chaîne de données ni des règles partagées.",
        "Les e-mails sont un canal complémentaire, jamais une barrière de sécurité. "
        "L'authentification locale est réservée à la démonstration ; la production "
        "requiert OIDC, RBAC et exploitation des journaux.",
        "Les modes AMDEC issus d'OCP conservent leur source et leur ligne ; les "
        "règles dérivées sont identifiées comme telles.",
    ):
        document.add_paragraph(f"• {text}")


def set_a4(document: Document) -> None:
    for section in document.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)


def main() -> int:
    if not REPORT.exists() or not METRICS.exists():
        raise FileNotFoundError("Rapport DOCX ou métriques finales absents")
    metrics = json.loads(METRICS.read_text(encoding="utf-8"))
    tests = metrics["tests"]["tests"]
    alerts = metrics["model"]["alert_hours_historical"]
    episodes = metrics["model"]["episodes"]
    replacements = {
        "111 tests automatisés": f"{tests} tests automatisés",
        "112 tests automatisés": f"{tests} tests automatisés",
        "119 tests automatisés": f"{tests} tests automatisés",
        "Tests automatisés | 119": f"Tests automatisés | {tests}",
        "1 662 heures atypiques": f"{alerts} heures atypiques",
        "1 589 heures atypiques": f"{alerts} heures atypiques",
        "1 662 points d’alarme": f"{alerts} points d’alarme",
        "1 589 points d’alarme": f"{alerts} points d’alarme",
        "68 épisodes": f"{episodes} épisodes",
        "70 épisodes": f"{episodes} épisodes",
        "Les 17 features du modèle": "Les 10 features contractuelles du modèle",
        "17 features": "10 features contractuelles",
        "28 points d'entrée": "42 routes /api/",
        "18 points d’entrée": "42 routes `/api/`",
        "faisceau de 221 tubes visibles": "faisceau illustratif non coté",
        "modèle WebGL vertical": "modèle WebGL horizontal conceptuel",
        "maintenance prédictive": "surveillance comportementale",
        "1 589 heures brutes": f"{alerts} heures brutes",
        "De 1 589 heures": f"De {alerts} heures",
        "8 274": "8 235",
        "19,2 %": "10,4 %",
        "facteur 23": "facteur 11,1",
        "17 grandeurs physiques": "10 features contractuelles ordonnées",
        "0,480": "0,973",
    }

    document = Document(REPORT)
    replace_text(document, replacements)
    synchronize_model_tables(document, metrics)
    append_addendum(document, metrics)
    style_tables(document)
    set_a4(document)
    document.core_properties.title = (
        "Surveillance comportementale du refroidisseur d’acide E7301"
    )
    document.save(REPORT)
    print(f"Rapport synchronisé : {REPORT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
